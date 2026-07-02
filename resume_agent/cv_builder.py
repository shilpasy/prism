"""
Generalized CV builder — no hardcoded contact info.
Builds the tailored .docx from a cv_data dict.
"""
from __future__ import annotations
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FONT_NAME = "Arial"
NAME_SIZE = Pt(16)
CONTACT_SIZE = Pt(9.5)
SECTION_SIZE = Pt(11)
ROLE_SIZE = Pt(10.5)
BODY_SIZE = Pt(10)
MARGIN = Inches(0.75)
NAVY = RGBColor(0x2E, 0x40, 0x57)
GRAY = RGBColor(0x70, 0x70, 0x70)
LINK_BLUE = RGBColor(0x1A, 0x6E, 0xBD)

EM_DASH_RE = re.compile(r"\s*[—–]\s*")


def sanitize(text: str) -> str:
    return EM_DASH_RE.sub(" | ", text)


def _apply_font(run, size, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _set_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN


def _add_border(paragraph, color="2E4057"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_header(doc: Document, title: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    _apply_font(run, SECTION_SIZE, bold=True, color=NAVY)
    _add_border(p)


def _bullet(doc: Document, text: str):
    pb = doc.add_paragraph(style="List Bullet")
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after = Pt(1)
    pb.paragraph_format.left_indent = Inches(0.25)
    _apply_font(pb.add_run(sanitize(text)), BODY_SIZE)


def build_docx(cv_data: dict, output_path: str) -> str:
    """Build a CV docx from cv_data. Returns path."""
    contact = cv_data["contact"]

    doc = Document()
    _set_margins(doc)
    doc.styles["Normal"].font.name = FONT_NAME
    doc.styles["Normal"].font.size = BODY_SIZE

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _apply_font(p.add_run(contact["name"]), NAME_SIZE, bold=True, color=NAVY)

    line1_parts = [contact["email"], contact["phone"]]
    line2_parts = [contact["linkedin"]]
    if contact.get("github"):
        line2_parts.append(contact["github"])
    if contact.get("scholar_url"):
        line2_parts.append(contact["scholar_url"])
    line2_parts.append(contact["location"])

    for line in ("  |  ".join(line1_parts), "  |  ".join(line2_parts)):
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before = Pt(0)
        pc.paragraph_format.space_after = Pt(1)
        _apply_font(pc.add_run(line), CONTACT_SIZE)

    # Summary (if present)
    if cv_data.get("summary"):
        _section_header(doc, "Professional Summary")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        _apply_font(p.add_run(sanitize(cv_data["summary"])), BODY_SIZE, italic=True)

    # Experience
    if cv_data.get("experiences"):
        _section_header(doc, "Experience")
        for exp in cv_data["experiences"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(f"{sanitize(exp['role'])}  |  {exp['org']}")
            _apply_font(r1, ROLE_SIZE, bold=True)
            if exp.get("dates"):
                _apply_font(p.add_run(f"  |  {exp['dates']}"), ROLE_SIZE, color=GRAY)
            for b in exp.get("bullets", []):
                _bullet(doc, b["text"])

    # Projects
    if cv_data.get("projects"):
        _section_header(doc, "Selected Projects")
        for proj in cv_data["projects"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(1)
            _apply_font(p.add_run(sanitize(proj["name"])), ROLE_SIZE, bold=True)
            for b in proj.get("bullets", []):
                _bullet(doc, b["text"])

    # Skills
    if cv_data.get("skills"):
        _section_header(doc, "Skills")
        for cat, items in cv_data["skills"].items():
            if not items:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _apply_font(p.add_run(f"{cat}: "), BODY_SIZE, bold=True)
            _apply_font(p.add_run(", ".join(str(i) for i in items)), BODY_SIZE)

    # Achievements
    if cv_data.get("achievements"):
        _section_header(doc, "Selected Achievements")
        for a in cv_data["achievements"]:
            _bullet(doc, a["text"])

    # Education
    if cv_data.get("education"):
        _section_header(doc, "Education")
        for edu in cv_data["education"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            r1 = p.add_run(f"{edu['degree']}  |  {edu['institution']}")
            _apply_font(r1, ROLE_SIZE, bold=True)
            trailing = "  |  ".join(filter(None, [edu.get("dates", ""), edu.get("location", "")]))
            if trailing:
                _apply_font(p.add_run(f"  |  {trailing}"), ROLE_SIZE, color=GRAY)
            for detail in edu.get("details", []):
                _bullet(doc, detail)

    doc.save(output_path)
    return output_path

